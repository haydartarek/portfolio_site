from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIRECTORY = Path(r"C:\Users\haydar\Desktop\cv_werkezoken")
DOCX_PATH = OUTPUT_DIRECTORY / "Haydar_Al-Egayli_Java_Developer_ATS_NL.docx"
PDF_PATH = OUTPUT_DIRECTORY / "Haydar_Al-Egayli_Java_Developer_ATS_NL.pdf"

NAVY = RGBColor(11, 39, 66)
BLUE = RGBColor(15, 76, 129)
TEXT = RGBColor(39, 54, 74)
DARK = RGBColor(17, 24, 39)
MUTED = RGBColor(51, 65, 85)


def set_cell_free_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.03

    rpr = normal.element.get_or_add_rPr()
    fonts = rpr.rFonts
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:cs"), "Arial")
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "nl-BE")
    lang.set(qn("w:eastAsia"), "nl-BE")
    lang.set(qn("w:bidi"), "nl-BE")
    rpr.append(lang)


def style_paragraph(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=0,
    space_after=3,
    keep_with_next=False,
    left_indent=0,
    first_line_indent=0,
):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.03
    fmt.keep_with_next = keep_with_next
    fmt.left_indent = Cm(left_indent)
    fmt.first_line_indent = Cm(first_line_indent)
    return paragraph


def set_run_style(run, *, size=10, bold=False, color=TEXT):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def add_centered_line(document, text, *, size, bold=False, color=DARK, after=1):
    paragraph = document.add_paragraph()
    style_paragraph(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=after,
    )
    set_run_style(
        paragraph.add_run(text),
        size=size,
        bold=bold,
        color=color,
    )
    return paragraph


def add_section_heading(document, text):
    paragraph = document.add_paragraph()
    style_paragraph(
        paragraph,
        space_before=7,
        space_after=3.5,
        keep_with_next=True,
    )
    set_run_style(
        paragraph.add_run(text.upper()),
        size=11.5,
        bold=True,
        color=BLUE,
    )
    return paragraph


def add_entry_heading(document, text, *, before=3.5):
    paragraph = document.add_paragraph()
    style_paragraph(
        paragraph,
        space_before=before,
        space_after=1,
        keep_with_next=True,
    )
    set_run_style(paragraph.add_run(text), size=10.2, bold=True, color=DARK)
    return paragraph


def add_body(document, text, *, after=3):
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, space_after=after)
    set_run_style(paragraph.add_run(text), size=10, color=TEXT)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    style_paragraph(
        paragraph,
        space_after=1.3,
        left_indent=0.45,
        first_line_indent=-0.25,
    )
    set_run_style(paragraph.add_run("\u2022 "), size=9.8, bold=True, color=BLUE)
    set_run_style(paragraph.add_run(text), size=9.8, color=TEXT)
    return paragraph


def add_skill_line(document, label, value):
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, space_after=1.5)
    set_run_style(paragraph.add_run(f"{label}: "), size=9.8, bold=True, color=DARK)
    set_run_style(paragraph.add_run(value), size=9.8, color=TEXT)
    return paragraph


def add_page_break(document):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    paragraph.paragraph_format.space_after = Pt(0)


def add_contact_block(document):
    add_centered_line(
        document,
        "HAYDAR AL-EGAYLI",
        size=19,
        bold=True,
        color=NAVY,
        after=0.5,
    )
    add_centered_line(
        document,
        "Junior Java Developer | Backend Developer",
        size=11.5,
        bold=True,
        color=BLUE,
        after=2.5,
    )
    add_centered_line(
        document,
        "2018 Antwerpen | +32 465 555 708 | heyder206@gmail.com | "
        "linkedin.com/in/haydartarek-dev | github.com/haydartarek",
        size=9,
        color=MUTED,
        after=0.8,
    )
    add_centered_line(
        document,
        "Rijbewijs B | Eigen wagen | Open voor functies in Antwerpen, Brussel en omgeving",
        size=9,
        color=MUTED,
        after=4,
    )


def add_profile(document):
    add_section_heading(document, "Professioneel profiel")
    add_body(
        document,
        "VDAB-gecertificeerde junior Java developer met een afgeronde opleiding "
        "Enterprise Java ontwikkelaar bij INTEC Brussel en praktijkervaring in het "
        "bouwen van full-stack toepassingen. Sterk in Java, Spring Boot, REST API's, "
        "Spring Security, Hibernate/JPA, SQL en gelaagde architectuur. Combineert "
        "softwareontwikkeling met eerdere ervaring in helpdesk, netwerkinfrastructuur "
        "en technische probleemoplossing. Zoekt een functie als Junior Java Developer "
        "of Backend Developer.",
        after=3.5,
    )


def add_skills(document):
    add_section_heading(document, "Technische vaardigheden")
    add_skill_line(document, "Programmeertalen", "Java, SQL, JavaScript, HTML5, CSS3")
    add_skill_line(
        document,
        "Java en backend",
        "Spring Framework, Spring Boot, Spring Security, JEE, REST API's, "
        "Hibernate/JPA, JDBC, Maven, Lombok",
    )
    add_skill_line(
        document,
        "Databanken",
        "MySQL, PostgreSQL, relationele datamodellering",
    )
    add_skill_line(
        document,
        "Frontend en web",
        "React, Angular, Node.js, responsive webdesign",
    )
    add_skill_line(
        document,
        "Testen en tools",
        "JUnit, Postman, Docker, Git, GitHub, UML, Active Directory",
    )
    add_skill_line(
        document,
        "Werkwijzen",
        "Objectgeoriënteerd programmeren, gelaagde architectuur, "
        "test-driven development, Agile/Scrum, CI/CD (basis)",
    )


def add_projects(document):
    add_section_heading(document, "Geselecteerde projecten")

    add_entry_heading(
        document,
        "ReadyRoad - Platform voor Belgische theorie-examens | "
        "Spring Boot, React, PostgreSQL, JWT",
    )
    add_bullet(
        document,
        "Ontwikkelt een full-stack platform voor de voorbereiding op het Belgische "
        "rijbewijs B-theorie-examen met een Spring Boot REST API en React-frontend.",
    )
    add_bullet(
        document,
        "Bouwde meertalige examenflows voor Nederlands, Engels, Frans en Arabisch "
        "met meer dan 202 vragen, JWT-authenticatie en rolgebaseerde toegang.",
    )
    add_bullet(document, "GitHub: github.com/haydartarek/readyroad_front_end")

    add_entry_heading(
        document,
        "Blog Central API Suite | Java, Spring Boot, Spring Security, "
        "Hibernate/JPA, MySQL",
    )
    add_bullet(
        document,
        "Ontwierp een modulaire backend API met een duidelijke "
        "controller-service-repositorystructuur.",
    )
    add_bullet(
        document,
        "Implementeerde beveiligde endpoints, validatie, foutafhandeling, "
        "contentbeheer en databaserelaties.",
    )
    add_bullet(document, "GitHub: github.com/haydartarek/blogSpring")

    add_entry_heading(
        document,
        "CoursePlatform - Leerplatform | Java, Spring Boot, Spring Security, "
        "REST API's, JWT",
    )
    add_bullet(
        document,
        "Bouwde een full-stack leerplatform met rolgebaseerd gebruikers- en "
        "cursusbeheer, Hibernate en responsieve frontend-integratie.",
    )
    add_bullet(
        document,
        "Ontwierp een gelaagde REST API voor duidelijke scheiding van "
        "verantwoordelijkheden en onderhoudbare code.",
    )
    add_bullet(document, "GitHub: github.com/haydartarek/courseplatform")


def add_education(document):
    add_section_heading(document, "Opleiding en certificering")
    add_entry_heading(
        document,
        "Enterprise Java ontwikkelaar | INTEC Brussel / VDAB | "
        "17 maart 2025 - 13 maart 2026",
    )
    add_bullet(
        document,
        "Officieel VDAB-competentierapport met 7 competenties beoordeeld als "
        "Professioneel: Spring Framework, ICT-applicaties ontwikkelen, "
        "objectgeoriënteerd programmeren, Java, programmeren in een specifieke "
        "computertaal, technische specificaties opstellen en Java-toepassingen "
        "ontwikkelen met een framework.",
    )
    add_bullet(
        document,
        "Curriculum: Java SE, Spring Boot, JEE, JPA/Hibernate, JDBC, Maven, JUnit, "
        "Docker, Git/GitHub, SQL/MySQL, REST, Angular, Node.js en Agile/Scrum.",
    )
    add_bullet(
        document,
        "Aanvullende competenties beoordeeld als Goed omvatten algoritmen, SQL, "
        "Angular, Web API, front-endontwikkeling, test-driven development, ORM, "
        "MVC en responsive design.",
    )
    add_bullet(
        document,
        "Enterprise Java Developer-certificaat behaald in maart 2026.",
    )


def add_experience(document):
    add_section_heading(document, "Relevante IT-ervaring")
    add_entry_heading(
        document,
        "Helpdeskmedewerker / IT Support | N-ID, Irak | 2013 - 2015",
    )
    add_bullet(
        document,
        "Verleende technische ondersteuning voor hardware-, software- en "
        "netwerkproblemen binnen de organisatie.",
    )
    add_bullet(
        document,
        "Beheerde gebruikersaccounts en toegangsrechten in Active Directory en "
        "registreerde, volgde en loste incidenten op via ticketsystemen.",
    )
    add_bullet(
        document,
        "Ondersteunde LAN/WAN-infrastructuur, routers en switches via remote "
        "troubleshooting.",
    )

    add_entry_heading(document, "IT-technicus | Scopesky, Irak | 2011 - 2013")
    add_bullet(
        document,
        "Installeerde en onderhield netwerkinfrastructuur, bekabeling, hardware, "
        "routers, switches en CCTV-systemen.",
    )
    add_bullet(
        document,
        "Diagnosticeerde netwerk- en hardwarestoringen om de operationele "
        "uitvaltijd te beperken.",
    )

    add_section_heading(document, "Aanvullende werkervaring")
    add_entry_heading(
        document,
        "Winkelmedewerker / Kassamedewerker | BVBA Gima, België | 2017 - 2022",
    )
    add_bullet(
        document,
        "Verleende klantendienst, verwerkte betalingen en beheerde voorraad in "
        "een drukke winkelomgeving.",
    )
    add_bullet(
        document,
        "Werkte nauwkeurig samen met collega's om dagelijkse operationele "
        "doelstellingen te behalen.",
    )


def add_languages(document):
    add_section_heading(document, "Professionele vaardigheden")
    add_skill_line(
        document,
        "Sterktes",
        "Probleemoplossing, teamwork, communicatie, klantgerichtheid, "
        "tijdbeheer, aanpassingsvermogen en verantwoordelijkheid",
    )

    add_section_heading(document, "Talen en mobiliteit")
    add_skill_line(
        document,
        "Talen",
        "Arabisch (moedertaal), Nederlands (gemiddeld), Engels (gemiddeld)",
    )
    add_skill_line(
        document,
        "Mobiliteit",
        "Antwerpen, rijbewijs B en eigen wagen",
    )


def build_document():
    document = Document()
    set_cell_free_document_defaults(document)

    properties = document.core_properties
    properties.title = "Haydar Al-Egayli - Junior Java Developer"
    properties.subject = "ATS-vriendelijk Nederlandstalig CV"
    properties.author = "Haydar Al-Egayli"
    properties.keywords = (
        "Java, Spring Boot, Spring Framework, Backend Developer, REST API, "
        "Hibernate, JPA, SQL, Junior Java Developer"
    )

    add_contact_block(document)
    add_profile(document)
    add_skills(document)
    add_projects(document)
    add_page_break(document)
    add_education(document)
    add_experience(document)
    add_languages(document)

    document.save(DOCX_PATH)
    return document


if __name__ == "__main__":
    build_document()
    print(f"DOCX={DOCX_PATH}")
